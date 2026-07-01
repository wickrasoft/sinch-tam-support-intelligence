#!/usr/bin/env python3
"""Combine the individual assignment Word docs into one deliverable.

Why this is not a naive docxcompose merge:
  - docxcompose collapses single-section documents into ONE section, keeping only
    the first (master) document's section properties. Every source doc here uses a
    "different first page" layout (`titlePg`): the cover shows the logo header and
    NO footer / no page number. A single-section merge would therefore make the
    covers of docs 2..N wrongly display the running footer + continuous page number.
  - To preserve each document's FIRST PAGE exactly, every doc must stay in its OWN
    section (with its own `titlePg`). We let docxcompose import the content/images/
    styles cleanly, then insert a real next-page SECTION BREAK (cloning the section
    properties, including `titlePg` + header/footer references) at each doc boundary.

Result: each document starts on a new page as its own section, with its original
first-page formatting intact. All source docs share the same corporate template,
so the cloned header/footer references resolve to the same logo header + footer.

Run:
    .venv/bin/python scripts/combine_docs.py
"""

import os
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

# Where the finalized individual docs live (and where the combined doc is written).
TAM_DIR = "/Users/kalwic/Documents/Resume/Kalum Wickramatunga/Kalum - Sinch - TAM"

# Order of the combined deliverable.
ORDER = [
    "Scenario-1-Major-Escalation-Nordea.docx",
    "Scenario-2-Team-Member-Mistake.docx",
    "Scenario-3-TAM-Sales-Conflict.docx",
    "Zendesk-Tooling-Improvement-Proposal.docx",
    "Part-A-AI-in-TAM-Work.docx",
    "Part-B-TAM-Support-Intelligence.docx",
]

OUTPUT_NAME = "Sinch-TAM-Assignment-Combined.docx"


def combine(source_dir=TAM_DIR, order=ORDER, output_name=OUTPUT_NAME):
    paths = [os.path.join(source_dir, f) for f in order]
    missing = [p for p in paths if not os.path.exists(p)]
    if missing:
        raise FileNotFoundError("Missing source docs:\n  " + "\n  ".join(missing))

    master = Document(paths[0])
    composer = Composer(master)
    body = master.element.body

    # First element of each appended doc -> where we close the previous section.
    boundaries = []
    for p in paths[1:]:
        idx = composer.append_index()      # index of trailing body sectPr (end of current content)
        composer.append(Document(p))
        boundaries.append(body[idx])       # first body element of the doc just appended

    # Template section properties (cover incl. titlePg + header/footer references).
    sect_pr = body.find(qn("w:sectPr"))

    def make_section_break():
        p = OxmlElement("w:p")
        p_pr = OxmlElement("w:pPr")
        sect = deepcopy(sect_pr)
        for t in sect.findall(qn("w:type")):   # remove explicit type -> default nextPage
            sect.remove(t)
        p_pr.append(sect)
        p.append(p_pr)
        return p

    # Close each preceding document as its own section, preserving its first page.
    for first_el in boundaries:
        first_el.addprevious(make_section_break())

    out_path = os.path.join(source_dir, output_name)
    composer.save(out_path)
    return out_path


def _verify(out_path):
    d = Document(out_path)
    imgs = sum(1 for r in d.part.rels.values() if "image" in r.reltype)
    print(f"Combined: paragraphs={len(d.paragraphs)} tables={len(d.tables)} "
          f"images={imgs} sections={len(d.sections)}")
    for i, s in enumerate(d.sections):
        tp = s._sectPr.find(qn("w:titlePg")) is not None
        print(f"  section {i}: titlePg={tp} start_type={s.start_type}")


if __name__ == "__main__":
    path = combine()
    print("Saved ->", path)
    _verify(path)
