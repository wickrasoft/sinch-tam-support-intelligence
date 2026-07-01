#!/usr/bin/env python3
"""Re-emit the existing Word deliverables into the Sinch standard Word template.

For each source .docx we open the standard template (which carries the Sinch
header logo, footer, fonts, theme colours, and named styles), clear its example
body, and rebuild the content using the template's own styles:
  Title / Subtitle, Heading 1-3, List Bullet (1-3), List Number, the branded
  "Sinch" table style, and Normal body copy.
"""
import os
import re
import shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import RGBColor, Pt, Inches

TEMPLATE = "/Users/kalwic/Downloads/Standard Word template with examples.docx"
TAM_DIR = "/Users/kalwic/Documents/Resume/Kalum Wickramatunga/Kalum - Sinch - TAM"
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DESKTOP = os.path.expanduser("~/Desktop")
ASSETS = os.path.join(ROOT, "scripts", "assets", "screenshots")

# Inline image markers emitted by the JS generators, e.g. "[[IMG:reply|caption]]".
# Every image file in scripts/assets/screenshots whose name starts with the key
# (reply, reply-1, reply-2, ...) is embedded at that point, so dropping in extra
# screenshots keeps all of them. Missing files are skipped silently.
IMG_MARKER_RE = re.compile(r"^\s*\[\[IMG:([^|\]]+?)\s*(?:\|\s*([^\]]*))?\]\]\s*$")
IMG_EXTS = (".png", ".jpg", ".jpeg", ".gif")
IMG_MAX_WIDTH_IN = 6.2

FILES = [
    "Part-A-AI-in-TAM-Work.docx",
    "Part-B-TAM-Support-Intelligence.docx",
    "Zendesk-Tooling-Improvement-Proposal.docx",
    "Scenario-3-TAM-Sales-Conflict.docx",
    "Scenario-1-Major-Escalation-Nordea.docx",
    "Scenario-2-Team-Member-Mistake.docx",
]

BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3", "List Bullet 4", "List Bullet 5"]
NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3", "List Number 4", "List Number 5"]
GREY = RGBColor(0x66, 0x66, 0x66)

# The template's built-in heading styles are very large (Title 52pt, Heading 1
# 32pt, Heading 2 24pt, Heading 3 16pt). Override them with a calmer, more
# professional hierarchy applied consistently to every document.
HEADING_STYLE_SIZES = {
    "Title": 28,
    "Subtitle": 14,
    "Heading 1": 18,
    "Heading 2": 14,
    "Heading 3": 12,
    "Heading 4": 11,
}

# With balanced heading sizes there is no longer any need to demote headings or
# to special-case the cover title; every document uses the natural levels above.
HEADING_DEMOTE = {}
TITLE_SIZE_PT = {}


def apply_heading_sizes(tpl):
    for name, pt in HEADING_STYLE_SIZES.items():
        try:
            tpl.styles[name].font.size = Pt(pt)
        except KeyError:
            pass


_DIVIDER_RE = re.compile(r"^[\s\u2014\u2013\-_]{6,}$")


def humanize(text):
    """Strip AI-tell punctuation: em/en dashes, semicolons, spaced slashes,
    arrows, and decorative divider lines. Keeps domain shorthand like P1/P2."""
    if not text:
        return text
    if _DIVIDER_RE.match(text):
        return ""
    # en-dash between digits -> hyphen (ranges); otherwise a comma
    text = re.sub(r"(?<=\d)\s*\u2013\s*(?=\d)", "-", text)
    text = re.sub(r"\s*\u2013\s*", ", ", text)
    # em-dash -> comma
    text = re.sub(r"\s*\u2014\s*", ", ", text)
    # arrows -> "to"
    text = re.sub(r"\s*\u2192\s*", " to ", text)
    # semicolons -> comma
    text = re.sub(r"\s*;\s*", ", ", text)
    # spaced slash (list separator) -> comma; leaves P1/P2, UI/UX, 30/60/90
    text = re.sub(r"\s+/\s+", ", ", text)
    # tidy whitespace and stray punctuation
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"(,\s*){2,}", ", ", text)
    text = re.sub(r",\s*\.", ".", text)
    return text.strip()


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def build_numfmt_map(doc):
    """Return a function numId,ilvl -> numFmt ('bullet','decimal',...)."""
    try:
        numbering = doc.part.numbering_part.element
    except Exception:
        return lambda n, l: "bullet"
    num_to_abs = {}
    for num in numbering.findall(qn("w:num")):
        nid = num.get(qn("w:numId"))
        absref = num.find(qn("w:abstractNumId"))
        if absref is not None:
            num_to_abs[nid] = absref.get(qn("w:val"))
    abs_fmt = {}
    for abn in numbering.findall(qn("w:abstractNum")):
        absid = abn.get(qn("w:abstractNumId"))
        lvls = {}
        for lvl in abn.findall(qn("w:lvl")):
            ilvl = lvl.get(qn("w:ilvl"))
            fmt = lvl.find(qn("w:numFmt"))
            lvls[ilvl] = fmt.get(qn("w:val")) if fmt is not None else "bullet"
        abs_fmt[absid] = lvls

    def lookup(numId, ilvl):
        absid = num_to_abs.get(str(numId))
        if absid is None:
            return "bullet"
        return abs_fmt.get(absid, {}).get(str(ilvl), "bullet")

    return lookup


def list_info(p):
    pPr = p._p.pPr
    if pPr is None or pPr.numPr is None:
        return None
    numPr = pPr.numPr
    numId = numPr.numId.val if numPr.numId is not None else None
    ilvl = numPr.ilvl.val if numPr.ilvl is not None else 0
    if numId is None:
        return None
    return numId, ilvl


def clear_body(tpl):
    body = tpl.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def add_para_runs(tpl, src_para, style, alignment=None, indent=None, color=None):
    para = tpl.add_paragraph(style=style)
    if alignment is not None:
        para.alignment = alignment
    if indent is not None:
        para.paragraph_format.left_indent = indent
    runs = src_para.runs if src_para is not None else []
    if not runs and src_para is not None and src_para.text:
        r = para.add_run(humanize(src_para.text))
        if color is not None:
            r.font.color.rgb = color
    for run in runs:
        r = para.add_run(humanize(run.text))
        r.bold = run.bold
        r.italic = run.italic
        if color is not None:
            r.font.color.rgb = color
    return para


def find_screenshots(key):
    """Return sorted image paths in ASSETS whose stem starts with key."""
    if not os.path.isdir(ASSETS):
        return []
    matches = []
    for name in os.listdir(ASSETS):
        stem, ext = os.path.splitext(name)
        if ext.lower() not in IMG_EXTS:
            continue
        if stem == key or stem.startswith(key + "-") or stem.startswith(key + "_"):
            matches.append(os.path.join(ASSETS, name))
    return sorted(matches)


def add_screenshots(tpl, key, caption):
    """Embed every screenshot matching key, each centred with an optional
    caption. Returns the number of images added."""
    paths = find_screenshots(key)
    for idx, path in enumerate(paths):
        pic_para = tpl.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        try:
            run.add_picture(path, width=Inches(IMG_MAX_WIDTH_IN))
        except Exception:
            continue
        if caption and idx == len(paths) - 1:
            cap = tpl.add_paragraph(style="Normal")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(humanize(caption))
            cr.italic = True
            cr.font.size = Pt(9)
            cr.font.color.rgb = GREY
    return len(paths)


# Table typography and colours. Larger, more readable than the template default.
TABLE_HEADER_PT = 12
TABLE_BODY_PT = 11
HEADER_FILL = "1B3A5B"          # Sinch-style deep blue header band
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_COLOR = "B7C0CC"


def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER_COLOR)
        borders.append(el)
    tblPr.append(borders)


def add_table(tpl, src_table):
    nrows = len(src_table.rows)
    ncols = len(src_table.columns)
    t = tpl.add_table(rows=nrows, cols=ncols)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(t)

    for ri, row in enumerate(src_table.rows):
        for ci in range(ncols):
            try:
                src_cell = row.cells[ci]
            except IndexError:
                continue
            cell = t.cell(ri, ci)
            cell.text = humanize(src_cell.text)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)
            is_header = ri == 0
            if is_header:
                _shade_cell(cell, HEADER_FILL)
            for par in cell.paragraphs:
                par.paragraph_format.space_after = Pt(2)
                par.paragraph_format.space_before = Pt(2)
                for run in par.runs:
                    run.font.size = Pt(TABLE_HEADER_PT if is_header else TABLE_BODY_PT)
                    if is_header:
                        run.bold = True
                        run.font.color.rgb = HEADER_TEXT


def demote_heading(name, demote):
    if not demote or not name.startswith("Heading"):
        return name
    try:
        lvl = int(name.split()[1])
    except (IndexError, ValueError):
        return name
    return f"Heading {min(lvl + demote, 9)}"


def convert(src_path, out_paths, demote=0, title_size_pt=None):
    src = Document(src_path)
    fmt_lookup = build_numfmt_map(src)
    tpl = Document(TEMPLATE)
    clear_body(tpl)
    apply_heading_sizes(tpl)

    blocks = list(iter_block_items(src))

    # --- collect the cover block (leading centred title paragraphs) ---
    i = 0
    cover = []
    while i < len(blocks):
        b = blocks[i]
        if isinstance(b, Table):
            break
        name = (b.style.name if b.style is not None else "") or ""
        if name.startswith("Heading"):
            break
        txt = b.text.strip()
        if txt and b.alignment != WD_ALIGN_PARAGRAPH.CENTER and list_info(b) is None:
            break
        if txt and b.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            cover.append(humanize(txt))
        i += 1

    if cover:
        title_para = tpl.add_paragraph(cover[0], style="Title")
        if title_size_pt:
            for run in title_para.runs:
                run.font.size = Pt(title_size_pt)
        if len(cover) > 1:
            tpl.add_paragraph(cover[1], style="Subtitle")
        for extra in cover[2:]:
            para = tpl.add_paragraph(style="Normal")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(extra)
            r.font.color.rgb = GREY

    # --- render the rest ---
    for b in blocks[i:]:
        if isinstance(b, Table):
            add_table(tpl, b)
            continue
        marker = IMG_MARKER_RE.match(b.text or "")
        if marker:
            add_screenshots(tpl, marker.group(1).strip(), (marker.group(2) or "").strip())
            continue
        li = list_info(b)
        if li is not None:
            numId, ilvl = li
            fmt = fmt_lookup(numId, ilvl)
            lvl = min(int(ilvl), 4)
            style = NUMBER_STYLES[lvl] if fmt not in ("bullet", None) else BULLET_STYLES[lvl]
            add_para_runs(tpl, b, style)
        else:
            name = (b.style.name if b.style is not None else "Normal") or "Normal"
            if name.startswith("Heading"):
                add_para_runs(tpl, b, demote_heading(name, demote))
            else:
                src_indent = b.paragraph_format.left_indent
                align = b.alignment
                # Justify normal body copy; leave centred lines and indented
                # quotes as they are.
                if align != WD_ALIGN_PARAGRAPH.CENTER and src_indent is None:
                    align = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_para_runs(tpl, b, "Normal", alignment=align, indent=src_indent)

    for out in out_paths:
        tpl.save(out)
    return len(blocks)


def main():
    for fname in FILES:
        src = os.path.join(ROOT, fname)
        if not os.path.exists(src):
            src_tam = os.path.join(TAM_DIR, fname)
            if os.path.exists(src_tam):
                src = src_tam
            else:
                print(f"SKIP (not found): {fname}")
                continue
        # read source first, then write outputs (safe even if same path)
        outs = [os.path.join(ROOT, fname), os.path.join(TAM_DIR, fname), os.path.join(DESKTOP, fname)]
        n = convert(src, outs, demote=HEADING_DEMOTE.get(fname, 0), title_size_pt=TITLE_SIZE_PT.get(fname))
        print(f"OK  {fname}  ({n} source blocks)  ->  root + TAM + Desktop")


if __name__ == "__main__":
    main()
